"""Uploaded document store for research context."""

from __future__ import annotations

import re
import uuid
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from app.config import get_settings


def _docs_dir() -> Path:
    path = get_settings().data_path / "uploads"
    path.mkdir(parents=True, exist_ok=True)
    return path


def save_upload(filename: str, content: bytes, content_type: str = "") -> dict[str, Any]:
    if not content:
        raise ValueError("Empty file")

    doc_id = str(uuid.uuid4())
    original = filename or "upload.txt"
    safe_name = re.sub(r"[^a-zA-Z0-9._-]+", "_", original)[:80] or "upload.bin"
    path = _docs_dir() / f"{doc_id}_{safe_name}"
    path.write_bytes(content)

    text = extract_text(path, content_type=content_type, original_name=original).strip()
    if not text:
        # Keep the job usable even if extraction is weak.
        text = (
            f"[Uploaded file: {original}]\n"
            "Text could not be fully extracted from this file. "
            "Use the filename and query context during research."
        )

    meta_path = _docs_dir() / f"{doc_id}.meta.txt"
    meta_path.write_text(text[:20000], encoding="utf-8")
    return {
        "id": doc_id,
        "filename": original,
        "path": str(path),
        "chars": len(text),
        "preview": text[:400],
    }


def get_document_text(doc_id: str) -> str:
    meta_path = _docs_dir() / f"{doc_id}.meta.txt"
    if meta_path.exists():
        return meta_path.read_text(encoding="utf-8")
    return ""


def get_documents_text(doc_ids: list[str]) -> str:
    chunks: list[str] = []
    for doc_id in doc_ids:
        text = get_document_text(doc_id).strip()
        if text:
            chunks.append(f"[Document {doc_id[:8]}]\n{text[:6000]}")
    return "\n\n".join(chunks)


def extract_text(path: Path, content_type: str = "", original_name: str = "") -> str:
    name = (original_name or path.name).lower()
    ctype = (content_type or "").lower()

    if name.endswith((".txt", ".md")) or ctype.startswith("text/"):
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""

    if name.endswith(".pdf") or "pdf" in ctype:
        return _extract_pdf(path)

    if name.endswith(".docx") or "wordprocessingml" in ctype or "officedocument" in ctype:
        return _extract_docx(path)

    if name.endswith(".doc"):
        # Old binary .doc is not supported cleanly; avoid crashing.
        return ""

    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        parts = [(page.extract_text() or "") for page in reader.pages[:30]]
        return "\n".join(parts).strip()
    except Exception:
        return ""


def _extract_docx(path: Path) -> str:
    # Prefer python-docx
    try:
        from docx import Document

        document = Document(str(path))
        parts: list[str] = []
        for para in document.paragraphs:
            if para.text and para.text.strip():
                parts.append(para.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
        if text:
            return text
    except Exception:
        pass

    # Fallback: raw document.xml
    try:
        with zipfile.ZipFile(path) as zf:
            xml = zf.read("word/document.xml")
        root = ElementTree.fromstring(xml)
        texts = [
            (node.text or "").strip()
            for node in root.iter()
            if node.tag.endswith("}t") and (node.text or "").strip()
        ]
        return "\n".join(texts).strip()
    except Exception:
        return ""
